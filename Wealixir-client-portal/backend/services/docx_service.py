"""DOCX generation service — stub for future feature."""

from __future__ import annotations


def generate_registration_docx(data: dict) -> bytes:  # type: ignore[type-arg]
    """Generate a .docx summary document for a registration.

    This is a stub. Install `python-docx` and implement as needed.

    Args:
        data: Registration data dict (same shape as pdf_service expects).

    Returns:
        Raw bytes of the generated .docx file.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from io import BytesIO

        doc = Document()
        doc.add_heading('Wealixir — Client Registration Summary', level=1)

        fields = [
            ('Full Name', data.get('name', '-')),
            ('Email Address', data.get('email', '-')),
            ('Mobile Number', data.get('mobile', '-')),
            ('Date of Birth', data.get('dob', '-')),
            ('PAN Number', data.get('pan_number', '-')),
            ('Address', data.get('address') or '-'),
            ('City', data.get('city') or '-'),
            ('Notes', data.get('notes') or '-'),
            ('Submitted At', data.get('submitted_at', '-')),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'

        for label, value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        raise NotImplementedError(
            'python-docx is not installed. '
            'Add it to requirements.txt and reinstall dependencies.'
        )
